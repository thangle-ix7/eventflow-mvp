from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = "docs"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(8)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    r = meta.add_run(subtitle)
    r.italic = True
    r.font.color.rgb = RGBColor.from_string("555555")


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F3A5F")


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], "F2F4F7")
        p = header_cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("0B2545")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(value))
    set_table_widths(table, widths)
    doc.add_paragraph()


def add_footer(doc, label):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(label)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("666666")


def save_doc(filename, title, subtitle, builder):
    doc = Document()
    configure_styles(doc)
    add_title(doc, title, subtitle)
    builder(doc)
    add_footer(doc, "EventFlow MVP - Tai lieu noi bo")
    doc.save(f"{OUTPUT_DIR}/{filename}")


def build_event_template(doc):
    add_note(doc, "Quyet dinh thiet ke: Template chi clone cau truc cong viec, khong tu tinh deadline cho task.")
    doc.add_heading("1. Muc Tieu", level=1)
    doc.add_paragraph(
        "Tinh nang Event Template giup Event Leader tao nhanh mot event moi tu bo khung co san. "
        "Template luu cau truc departments, tasks, subtasks va category; event moi sau khi clone se doc lap hoan toan voi template goc."
    )

    doc.add_heading("2. Nguyen Tac Nghiep Vu", level=1)
    add_bullets(doc, [
        "Template la mot ban ghi Event co nature = TEMPLATE.",
        "Event van hanh that la Event co nature = NORMAL.",
        "Khong tao cac bang rieng nhu template_tasks hay template_departments trong MVP.",
        "Task trong template duoc phep khong co deadline.",
        "Khi instantiate, deadline, assignee va progress khong clone tu template.",
        "Sau khi clone, leader vao man hinh hoan thien ke hoach de gan deadline va assignee theo thuc te.",
    ])

    doc.add_heading("3. Entity Va Database Can Sua", level=1)
    add_table(doc, ["Thanh phan", "Thay doi", "Ghi chu"], [
        ["events", "Them nature VARCHAR(20) DEFAULT 'NORMAL'", "Phan biet TEMPLATE va NORMAL."],
        ["tasks", "Cho phep deadline NULL", "Can sua validation de template task khong bat buoc deadline."],
        ["task_categories", "Bang moi: id, event_id, name, description, color", "Dung de phan loai task trong template va event that."],
        ["tasks", "Them category_id nullable", "Gan task voi category neu co."],
        ["Event.java", "Them field nature", "Co the dung enum EventNature hoac String co validate."],
        ["Task.java", "deadline nullable, them category", "Khong them deadlineOffsetDays trong scope hien tai."],
    ], [1900, 3700, 3760])

    doc.add_heading("4. Du Lieu Duoc Clone", level=1)
    add_table(doc, ["Clone", "Khong clone / reset"], [
        ["Event name/description co the lay tu template hoac request moi", "nature: chuyen thanh NORMAL"],
        ["Departments: name, description, leader optional neu muon", "createdAt sinh moi"],
        ["Task: title, description, priority, category, parent/subtask structure", "deadline = null, assignee = null, status = TODO, progressPercentage = 0"],
        ["Task category: name, description, color", "Khong lien ket nguoc voi template sau khi clone"],
    ], [4680, 4680])

    doc.add_heading("5. API Can Xay Dung", level=1)
    add_table(doc, ["Method", "Endpoint", "Muc dich", "Quyen"], [
        ["GET", "/api/v1/events?nature=TEMPLATE", "Lay danh sach template user duoc xem", "Member cua event/template"],
        ["GET", "/api/v1/events/{templateId}", "Preview template", "Member cua event/template"],
        ["POST", "/api/v1/events", "Tao event NORMAL hoac TEMPLATE", "Event Leader"],
        ["POST", "/api/v1/events/{templateId}/instantiate", "Clone template thanh event moi", "Event Leader"],
        ["PUT", "/api/v1/events/{templateId}", "Sua metadata template", "Event Leader"],
    ], [1000, 3100, 3350, 1910])

    doc.add_heading("6. Request Instantiate De Xuat", level=1)
    doc.add_paragraph(
        '{ "name": "Workshop AI Thang 7", "startTime": "2026-07-01T08:00:00", '
        '"endTime": "2026-07-31T17:00:00", "location": "HCM" }'
    )

    doc.add_heading("7. Checklist Trien Khai", level=1)
    add_numbered(doc, [
        "Tao migration cho events.nature, tasks.deadline nullable, task_categories va tasks.category_id.",
        "Cap nhat Entity, DTO va validation theo nature cua event.",
        "Them filter nature vao EventRepository/EventService.",
        "Xay TemplateInstantiationService clone event, departments, categories, tasks va subtasks.",
        "Them API instantiate va test case clone doc lap.",
        "Cap nhat frontend luong chon template va man hinh hoan thien task sau clone.",
    ])


def build_ai_suggestion(doc):
    add_note(doc, "Quyet dinh thiet ke: AI chi dua ra goi y. User xem, chinh sua va xac nhan truoc khi luu vao database.")
    doc.add_heading("1. Muc Tieu", level=1)
    doc.add_paragraph(
        "AI Suggestion ho tro Event Leader va Department Leader tao nhanh department, planning, milestone, task va subtask. "
        "Backend can lay du lieu hien co cua event, tao prompt co cau truc, nhan JSON tu AI provider va tra ve DTO an toan cho frontend."
    )

    doc.add_heading("2. Pham Vi Da Co Va Can Bo Sung", level=1)
    add_table(doc, ["Hang muc", "Trang thai hien tai", "Can lam tiep"], [
        ["Department suggestion", "Da co API", "Bo sung context objective/contextDescription/scale neu them vao Event."],
        ["Task suggestion", "Da co API theo event", "Bo sung API suggest task theo milestone."],
        ["Calendar suggestion", "Da co API", "Giu nguyen, khong nam trong scope tai lieu moi."],
        ["Subtask suggestion", "Da co API", "Bo sung context milestone neu task gan milestone."],
        ["Planning suggestion", "Chua co", "Them entity, DTO, service va API."],
        ["Milestone suggestion", "Chua co", "Them entity, DTO, service va API."],
    ], [2100, 2700, 4560])

    doc.add_heading("3. Entity Va Database Can Them", level=1)
    add_table(doc, ["Entity", "Field chinh", "Quan he"], [
        ["Planning", "id, event_id, title, description, created_by, created_at, updated_at", "Mot event co nhieu planning."],
        ["PlanningPhase", "id, planning_id, phase_name, description, objective, order_index", "Mot planning co nhieu phase."],
        ["Milestone", "id, event_id, name, description, expected_deadline, expected_result, priority, status", "Milestone thuoc event va doc lap voi planning."],
        ["Task", "milestone_id nullable", "Task co the thuoc mot milestone."],
        ["Event", "context_description, objective, expected_attendees/scale optional", "Tang chat luong prompt AI."],
    ], [1700, 5000, 2660])

    doc.add_heading("4. API AI Suggestion", level=1)
    add_table(doc, ["Method", "Endpoint", "Muc dich", "Quyen"], [
        ["POST", "/api/v1/ai-suggestions/events/{eventId}/departments", "Goi y department", "Event Leader"],
        ["POST", "/api/v1/ai-suggestions/events/{eventId}/planning", "Goi y planning tong the", "Event Leader"],
        ["POST", "/api/v1/ai-suggestions/events/{eventId}/milestones", "Goi y milestone", "Event Leader"],
        ["POST", "/api/v1/ai-suggestions/events/{eventId}/tasks", "Goi y task theo event", "Event Leader, Department Leader"],
        ["POST", "/api/v1/ai-suggestions/milestones/{milestoneId}/tasks", "Goi y task theo milestone", "Event Leader, Department Leader"],
        ["POST", "/api/v1/ai-suggestions/tasks/{taskId}/subtasks", "Goi y subtask", "Event Leader, Department Leader"],
    ], [900, 3450, 3000, 2010])

    doc.add_heading("5. API Luu Du Lieu Sau Khi User Xac Nhan", level=1)
    add_table(doc, ["Method", "Endpoint", "Muc dich"], [
        ["GET/POST", "/api/v1/events/{eventId}/plannings", "Lay danh sach hoac tao planning."],
        ["GET/PUT/DELETE", "/api/v1/plannings/{planningId}", "Chi tiet, cap nhat hoac xoa planning."],
        ["POST", "/api/v1/plannings/{planningId}/phases", "Them phase vao planning."],
        ["PUT/DELETE", "/api/v1/planning-phases/{phaseId}", "Cap nhat hoac xoa phase."],
        ["GET/POST", "/api/v1/events/{eventId}/milestones", "Lay danh sach hoac tao milestone."],
        ["GET/PUT/DELETE", "/api/v1/milestones/{milestoneId}", "Chi tiet, cap nhat hoac xoa milestone."],
    ], [1300, 3700, 4360])

    doc.add_heading("6. Luong Xu Ly Backend", level=1)
    add_numbered(doc, [
        "Kiem tra quyen cua user theo event hoac department.",
        "Load event, departments, planning, phases, milestones va tasks lien quan.",
        "Tao prompt yeu cau AI tra JSON object, khong markdown.",
        "Parse JSON, validate enum, gioi han do dai text va loc id khong hop le.",
        "Tra suggestion DTO ve frontend, chua ghi database neu user chua xac nhan.",
        "Khi user xac nhan, goi API CRUD tuong ung de luu du lieu.",
    ])

    doc.add_heading("7. Checklist Trien Khai", level=1)
    add_bullets(doc, [
        "Them migration va entity cho Planning, PlanningPhase, Milestone.",
        "Them DTO response/request cho planning, phase, milestone va AI suggestion.",
        "Them repository va service CRUD.",
        "Mo rong AiSuggestionService voi planningInstructions va milestoneInstructions.",
        "Viet test phan quyen: Member khong duoc goi AI suggestion.",
        "Cap nhat frontend flow: goi y -> preview -> edit -> confirm save.",
    ])


def build_workload(doc):
    add_note(doc, "Quyet dinh thiet ke MVP: Workload Score tinh truc tiep khi goi API, khong luu score vao tasks hay summary table.")
    doc.add_heading("1. Muc Tieu", level=1)
    doc.add_paragraph(
        "Workload Score giup Event Leader va Department Leader nhin nhanh muc do tai cong viec cua tung thanh vien. "
        "Chi so nay duoc tinh dua tren so task chua hoan thanh dang assign cho member so voi suc chua cong viec toi da."
    )

    doc.add_heading("2. Cong Thuc Va Trang Thai", level=1)
    doc.add_paragraph("workloadScore = memberAssignedTasks / teamAverageAssignedTasks * 100")
    add_table(doc, ["Khoang diem", "Status", "Y nghia"], [
        ["0% - 70%", "LOW", "Dang nhan it task hon mat bang chung."],
        ["71% - 120%", "NORMAL", "Khoi luong cong viec gan voi muc trung binh."],
        ["121% - 160%", "HIGH", "Dang nhan nhieu task hon trung binh, can theo doi."],
        ["> 160%", "OVERLOADED", "Lech tai ro ret, leader nen can bang lai."],
    ], [1800, 1900, 5660])

    doc.add_heading("3. Entity Va Database Can Sua", level=1)
    add_table(doc, ["Thanh phan", "Thay doi", "Ghi chu"], [
        ["event_members", "Khong them field workload", "Workload khong phai gioi han giao viec."],
        ["EventMember.java", "Khong them maxTaskCapacity", "Score duoc tinh runtime tu task dang assign."],
        ["WorkloadStatus", "Enum/DTO: LOW, NORMAL, HIGH, OVERLOADED", "Khong nhat thiet luu database trong MVP."],
        ["tasks", "Khong them workload_score", "Score la chi so tong hop theo member, khong thuoc task."],
        ["team_workload_summary", "Chua can tao", "Chi can xem xet sau khi du lieu lon."],
    ], [2100, 3500, 3760])

    doc.add_heading("4. Query Tinh Workload", level=1)
    add_bullets(doc, [
        "currentAssignedTasks: dem tasks co assignee_id = userId va status != DONE.",
        "completedTasks: dem tasks co assignee_id = userId va status = DONE.",
        "inProgressTasks: dem tasks co assignee_id = userId va status khac DONE.",
        "teamAverageAssignedTasks: tong task dang assign trong scope chia cho so member trong scope.",
        "Nen tinh theo eventId de tranh cong don task cua event khac.",
        "Neu departmentId duoc truyen, chi tinh task trong department do.",
    ])

    doc.add_heading("5. API Can Xay Dung", level=1)
    add_table(doc, ["Method", "Endpoint", "Muc dich", "Quyen"], [
        ["GET", "/api/v1/events/{eventId}/departments/{departmentId}/workload", "Dashboard workload cua department", "Event Leader, Department Leader"],
        ["GET", "/api/v1/events/{eventId}/members/{userId}/workload", "Chi tiet workload member", "Leader hoac chinh member"],
        ["GET", "/api/v1/events/{eventId}/workload", "Tong quan workload toan event", "Event Leader"],
    ], [900, 3750, 2800, 1910])

    doc.add_heading("6. Response DTO De Xuat", level=1)
    add_table(doc, ["DTO", "Field chinh"], [
        ["MemberWorkloadResponse", "memberId, memberName, departmentId, assignedTasks, completedTasks, inProgressTasks, teamAverageAssignedTasks, workloadScore, workloadStatus, tasks"],
        ["DepartmentWorkloadResponse", "eventId, departmentId, departmentName, totalMembers, totalAssignedTasks, averageWorkloadScore, overloadedMembers, members"],
        ["EventWorkloadResponse", "eventId, totalDepartments, totalMembers, totalAssignedTasks, averageWorkloadScore, overloadedDepartmentCount, departments"],
    ], [2500, 6860])

    doc.add_heading("7. Phan Quyen", level=1)
    add_table(doc, ["Vai tro", "Duoc phep"], [
        ["Event Leader", "Xem workload toan event, department va member."],
        ["Department Leader", "Xem workload department minh quan ly va member trong department."],
        ["Member", "Chi xem workload ca nhan va task duoc giao."],
    ], [2200, 7160])

    doc.add_heading("8. Checklist Trien Khai", level=1)
    add_numbered(doc, [
        "Khong tao migration capacity hoac hard limit workload.",
        "Them query count task theo eventId, assigneeId, departmentId va status.",
        "Xay WorkloadService tinh score va status runtime.",
        "Them WorkloadController voi cac endpoint doc dashboard/detail.",
        "Viet test phan quyen va test cong thuc status.",
        "Cap nhat frontend dashboard/canh bao qua tai.",
    ])


save_doc(
    "event-template-solution.docx",
    "Event Template Solution",
    "Tài liệu tách biệt cho team backend/frontend - EventFlow MVP",
    build_event_template,
)
save_doc(
    "ai-suggestion-solution.docx",
    "AI Suggestion Solution",
    "Tài liệu tách biệt cho team backend/frontend - EventFlow MVP",
    build_ai_suggestion,
)
save_doc(
    "workload-score-solution.docx",
    "Workload Score Solution",
    "Tài liệu tách biệt cho team backend/frontend - EventFlow MVP",
    build_workload,
)
