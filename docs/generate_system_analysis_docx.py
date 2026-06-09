from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/EventFlow_Tai_lieu_phan_tich_he_thong.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 33, 36)
MUTED = RGBColor(90, 96, 106)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9DEE7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    run.font.size = Pt(9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def paragraph(doc, text="", style=None, bold_lead=None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        run = p.add_run(bold_lead)
        run.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def bullet(doc, text):
    return paragraph(doc, text, style="List Bullet")


def number(doc, text):
    return paragraph(doc, text, style="List Number")


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [9360])
    set_table_borders(table, "D6DCE6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p.add_run(" " + text)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        r.bold = True
        r.font.color.rgb = DARK_BLUE
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx else WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(value)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()
    return table


def add_doc_header_footer(doc):
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "EventFlow | Tài liệu phân tích hệ thống"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)
    for run in footer_p.runs:
        run.font.color.rgb = MUTED


def build_document():
    doc = Document()
    setup_styles(doc)
    add_doc_header_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Tài liệu phân tích hệ thống EventFlow")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = DARK_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Tổng quan nghiệp vụ, luồng sử dụng, phân quyền và giá trị hệ thống")
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(meta, [2400, 5200])
    set_table_borders(meta, "E1E5ED")
    meta_rows = [
        ("Hệ thống", "EventFlow MVP"),
        ("Phạm vi", "Quản lý điều phối công việc nội bộ cho quá trình tổ chức sự kiện"),
        ("Công nghệ chính", "React/Vite, Spring Boot, PostgreSQL, Docker Compose, JWT"),
        ("Đối tượng đọc", "Nhóm phát triển, giảng viên/hội đồng, khách hàng nghiệp vụ và đội vận hành"),
    ]
    for row, (label, value) in zip(meta.rows, meta_rows):
        set_cell_shading(row.cells[0], LIGHT_FILL)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)
        for cell in row.cells:
            set_cell_margins(cell, top=100, bottom=100)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)

    doc.add_page_break()

    heading(doc, "Mục lục", 1)
    toc_items = [
        "1. Giới thiệu tổng quan hệ thống",
        "2. Vấn đề thực tế cần giải quyết",
        "3. Đối tượng sử dụng hệ thống",
        "4. Danh sách nhóm tính năng chính",
        "5. Mô tả chi tiết từng tính năng theo nghiệp vụ",
        "6. Luồng sử dụng hệ thống",
        "7. Phân quyền người dùng",
        "8. Giá trị hệ thống mang lại",
        "9. Kết luận",
    ]
    for item in toc_items:
        bullet(doc, item)

    callout(
        doc,
        "Tóm tắt:",
        "EventFlow là hệ thống quản lý điều phối sự kiện tập trung vào phân công, theo dõi tiến độ, báo cáo, tài liệu, dashboard và thông báo nhắc việc cho ban tổ chức.",
    )

    heading(doc, "1. Giới thiệu tổng quan hệ thống", 1)
    paragraph(
        doc,
        "EventFlow là hệ thống hỗ trợ lập kế hoạch, phân công và theo dõi tiến độ công việc trong quá trình tổ chức sự kiện. Hệ thống tập trung vào bài toán quản lý nội bộ của ban tổ chức: tạo sự kiện, chia phòng ban, thêm thành viên, giao việc, cập nhật tiến độ, nộp báo cáo, đính kèm tài liệu, theo dõi dashboard và nhận thông báo nhắc việc.",
    )
    paragraph(
        doc,
        "Trong phạm vi MVP, EventFlow phù hợp với câu lạc bộ, nhóm dự án, doanh nghiệp nhỏ hoặc đội tổ chức sự kiện cần điều phối theo mô hình leader và member. Hệ thống chưa tập trung vào bán vé, check-in khách tham dự hay thanh toán, mà ưu tiên điều phối công việc nội bộ trước, trong và sau sự kiện.",
    )
    heading(doc, "Thành phần kỹ thuật chính", 2)
    for item in [
        "Frontend: React, Vite, React Router, TanStack Query và Tailwind CSS.",
        "Backend: Spring Boot, Spring Security, Spring Data JPA, Flyway migration và OpenAPI/Swagger.",
        "Cơ sở dữ liệu: PostgreSQL.",
        "Xác thực: JWT, đăng ký tài khoản, xác minh email, đăng nhập, quên mật khẩu và đặt lại mật khẩu.",
        "Tích hợp: SMTP email, Telegram bot, AI assistant tương thích OpenAI/Groq, lưu trữ file local hoặc Supabase Storage/S3.",
        "Triển khai: Docker Compose với các service database, backend và frontend.",
    ]:
        bullet(doc, item)

    heading(doc, "2. Vấn đề thực tế cần giải quyết", 1)
    for item in [
        "Thông tin sự kiện, phòng ban, người phụ trách và deadline bị phân tán trên nhiều kênh như nhóm chat, bảng tính, email hoặc ghi chú cá nhân.",
        "Leader khó nắm tiến độ tổng thể, đặc biệt là việc nào quá hạn, việc nào đang chờ review và thành viên nào đang quá tải.",
        "Member thiếu một nơi thống nhất để xem việc được giao, cập nhật tiến độ, nộp minh chứng và nhận phản hồi.",
        "Tài liệu, ảnh báo cáo và link tham khảo nằm rải rác, khó tìm lại theo sự kiện, phòng ban hoặc task.",
        "Nhắc việc thủ công tốn thời gian và dễ bỏ sót các task sắp đến hạn hoặc đã quá hạn.",
        "Phân quyền không rõ ràng khiến dữ liệu có thể bị lộ quá phạm vi cần thiết hoặc người không phù hợp có thể thao tác nhầm.",
        "Khi sự kiện có nhiều phòng ban, việc tổng hợp báo cáo và đánh giá hiệu quả giữa các nhóm thường mất nhiều công sức.",
    ]:
        bullet(doc, item)
    paragraph(
        doc,
        "EventFlow giải quyết các vấn đề trên bằng cách đưa dữ liệu sự kiện, phòng ban, thành viên, task, báo cáo, file đính kèm, lịch và thông báo về một hệ thống tập trung có phân quyền.",
    )

    heading(doc, "3. Đối tượng sử dụng hệ thống", 1)
    heading(doc, "3.1. Leader sự kiện", 2)
    paragraph(doc, "Leader là người tạo và điều phối sự kiện. Nhóm này cần quan sát toàn cảnh, chia phòng ban, thêm thành viên, giao việc, điều chỉnh deadline, review kết quả và xem dashboard tiến độ.")
    for item in [
        "Tạo, cập nhật và xóa sự kiện.",
        "Thêm/xóa thành viên, thay đổi vai trò và gán thành viên vào phòng ban.",
        "Tạo phòng ban, cập nhật thông tin phòng ban và chỉ định người phụ trách.",
        "Tạo task/subtask, giao người thực hiện, đặt deadline, mức ưu tiên và trạng thái.",
        "Theo dõi dashboard theo sự kiện và phòng ban.",
        "Review báo cáo, xem tài liệu, lịch sử cập nhật và minh chứng.",
    ]:
        bullet(doc, item)
    heading(doc, "3.2. Member sự kiện", 2)
    paragraph(doc, "Member là người tham gia thực hiện công việc trong sự kiện. Nhóm này chủ yếu xem các công việc được giao, cập nhật tiến độ, nộp báo cáo và nhận thông báo.")
    for item in [
        "Xem danh sách sự kiện mình tham gia.",
        "Xem task được giao và thông tin phòng ban liên quan.",
        "Cập nhật trạng thái, phần trăm tiến độ và nội dung thực hiện.",
        "Tải lên minh chứng, ảnh báo cáo hoặc tài liệu liên quan đến task.",
        "Nhận thông báo trong ứng dụng, qua email hoặc Telegram nếu đã liên kết.",
    ]:
        bullet(doc, item)
    heading(doc, "3.3. Hệ thống tích hợp/phụ trợ", 2)
    for item in [
        "Email service gửi email xác minh tài khoản và đặt lại mật khẩu.",
        "Telegram bot liên kết tài khoản và gửi thông báo nhắc việc.",
        "Scheduler backend tạo và xử lý thông báo cho task sắp đến hạn, quá hạn hoặc workflow liên quan.",
        "AI assistant hỗ trợ người dùng thao tác, gợi ý task/action dựa trên ngữ cảnh trang hiện tại.",
    ]:
        bullet(doc, item)

    heading(doc, "4. Danh sách nhóm tính năng chính", 1)
    features = [
        "Xác thực và bảo mật tài khoản.",
        "Quản lý sự kiện.",
        "Quản lý thành viên và phân quyền theo sự kiện.",
        "Quản lý phòng ban.",
        "Quản lý công việc, subtask, trạng thái và tiến độ.",
        "Báo cáo công việc và review kết quả.",
        "Quản lý tài liệu/đính kèm.",
        "Dashboard và thống kê tiến độ.",
        "Lịch sự kiện, tài liệu tổng hợp và báo cáo theo sự kiện.",
        "Thông báo, nhắc việc và tích hợp Telegram/email.",
        "Hồ sơ cá nhân và tùy chỉnh trải nghiệm.",
        "AI assistant.",
        "Vận hành, audit log, rate limiting và bảo vệ API.",
    ]
    for item in features:
        number(doc, item)

    heading(doc, "5. Mô tả chi tiết từng tính năng theo nghiệp vụ", 1)
    detail_sections = [
        ("5.1. Xác thực và bảo mật tài khoản", [
            "Người dùng đăng ký tài khoản bằng tên, email và mật khẩu. Sau khi đăng ký, hệ thống gửi token xác minh email. Người dùng có thể đăng nhập sau khi tài khoản hợp lệ.",
            "Hệ thống hỗ trợ quên mật khẩu, gửi token đặt lại mật khẩu và tạo mật khẩu mới. Mật khẩu được mã hóa bằng BCrypt, API nội bộ sử dụng JWT bearer token.",
            "Các lớp bảo vệ gồm rate limiting, abuse protection theo độ dài request và audit log cho hành vi đáng chú ý.",
        ]),
        ("5.2. Quản lý sự kiện", [
            "Leader tạo sự kiện mới với tên, mô tả, địa điểm, thời gian bắt đầu, thời gian kết thúc và trạng thái.",
            "Người tạo sự kiện trở thành leader của sự kiện đó. Người dùng chỉ xem được các sự kiện mà mình là thành viên.",
            "Leader có quyền cập nhật và xóa sự kiện; member chỉ xem thông tin trong phạm vi sự kiện mình tham gia.",
        ]),
        ("5.3. Quản lý thành viên và vai trò", [
            "Vai trò được gán theo từng sự kiện, gồm LEADER và MEMBER. Một người có thể là leader ở sự kiện này nhưng là member ở sự kiện khác.",
            "Leader có thể thêm thành viên, thay đổi vai trò và xóa thành viên. Member được xem thông tin thành viên trong phạm vi hệ thống cho phép.",
        ]),
        ("5.4. Quản lý phòng ban", [
            "Phòng ban là đơn vị tổ chức trong sự kiện, ví dụ Hậu cần, Truyền thông, Nội dung, Kỹ thuật hoặc Đối ngoại.",
            "Leader có thể tạo, cập nhật, xóa phòng ban, gán thành viên vào phòng ban và chỉ định người phụ trách.",
            "Member có thể xem phòng ban liên quan, danh sách thành viên phòng ban trong phạm vi được phép và task được giao.",
        ]),
        ("5.5. Quản lý công việc và subtask", [
            "Task thuộc về một sự kiện, có thể thuộc phòng ban, có người được giao, tiêu đề, mô tả, deadline, trạng thái, mức ưu tiên và phần trăm tiến độ.",
            "Trạng thái task gồm TODO, IN_PROGRESS, IN_REVIEW và DONE. Mức ưu tiên gồm LOW, MEDIUM, HIGH và URGENT.",
            "Leader tạo, sửa, xóa task/subtask và thay đổi người được giao. Member chỉ xem và cập nhật các task được giao.",
        ]),
        ("5.6. Cập nhật tiến độ công việc", [
            "Member được giao task có thể cập nhật nội dung thực hiện và phần trăm tiến độ.",
            "Thông tin tiến độ giúp leader biết task đang ở mức nào, có nguy cơ trễ hạn hay cần hỗ trợ hay không.",
        ]),
        ("5.7. Báo cáo công việc và review", [
            "Member nộp báo cáo gồm phần trăm tiến độ, mô tả và ảnh minh chứng nếu có.",
            "Leader xem báo cáo, file đính kèm và lịch sử cập nhật để review. Khi review, hệ thống lưu trạng thái trước/sau và feedback.",
            "Luồng này phù hợp với nghiệp vụ nộp kết quả, chờ duyệt, chấp nhận hoàn thành hoặc yêu cầu bổ sung.",
        ]),
        ("5.8. Quản lý tài liệu và file đính kèm", [
            "Task có thể có file đính kèm hoặc link tài liệu. File lưu thông tin người tải lên, tên gốc, content type, kích thước, nhà cung cấp lưu trữ và đường dẫn.",
            "Attachment có ba mức hiển thị: TASK_ONLY, DEPARTMENT và EVENT_PUBLIC.",
            "Leader và assignee của task có thể tải lên file/link. Quyền xem file được kiểm soát theo quyền xem task và visibility.",
        ]),
        ("5.9. Dashboard và thống kê", [
            "Dashboard dành cho leader, gồm tổng quan task, tiến độ theo thời gian, task theo phòng ban, task theo assignee và task theo trạng thái.",
            "Hệ thống hỗ trợ dashboard cấp sự kiện, cấp phòng ban và so sánh chỉ số theo giai đoạn thời gian.",
            "Dashboard giúp leader phát hiện phòng ban chậm tiến độ, người quá tải và task cần ưu tiên xử lý.",
        ]),
        ("5.10. Lịch sự kiện", [
            "Member của sự kiện có thể xem lịch theo tháng. Leader có thể tạo và cập nhật calendar item.",
            "Lịch phù hợp để quản lý các mốc họp ban tổ chức, hạn chót nộp nội dung, ngày tổng duyệt, ngày diễn ra sự kiện và các mốc nhắc việc quan trọng.",
        ]),
        ("5.11. Thông báo và tích hợp", [
            "Hệ thống lưu thông báo với người nhận, task/calendar liên quan, kênh gửi, loại thông báo, trạng thái gửi, số lần retry, tiêu đề, nội dung và thời gian đọc.",
            "Các kênh gồm in-app notification, email và Telegram. Người dùng có thể xem thông báo, xem số chưa đọc, đánh dấu từng thông báo hoặc tất cả là đã đọc.",
        ]),
        ("5.12. Hồ sơ cá nhân và AI assistant", [
            "Người dùng có thể xem hồ sơ, cập nhật tùy chọn cá nhân, tải avatar và tạo token liên kết Telegram. Người dùng chỉ thao tác với hồ sơ của chính mình.",
            "AI assistant cho phép chat theo ngữ cảnh, hỗ trợ trả lời, gợi ý action hoặc tạo bản nháp task/action.",
        ]),
    ]
    for title, paragraphs in detail_sections:
        heading(doc, title, 2)
        for text in paragraphs:
            paragraph(doc, text)

    heading(doc, "6. Luồng sử dụng hệ thống", 1)
    flows = [
        ("6.1. Khởi tạo tài khoản", [
            "Người dùng đăng ký tài khoản bằng email và mật khẩu.",
            "Hệ thống gửi email xác minh.",
            "Người dùng xác minh email.",
            "Hệ thống kích hoạt tài khoản và cấp JWT sau khi đăng nhập thành công.",
            "Người dùng vào trang danh sách sự kiện.",
        ]),
        ("6.2. Tạo và cấu hình sự kiện", [
            "Leader tạo sự kiện mới.",
            "Hệ thống gán người tạo làm leader của sự kiện.",
            "Leader cập nhật mô tả, địa điểm, thời gian và trạng thái.",
            "Leader tạo phòng ban, thêm thành viên và gán thành viên vào phòng ban.",
        ]),
        ("6.3. Phân công công việc", [
            "Leader mở trang tasks của sự kiện.",
            "Leader tạo task với thông tin nghiệp vụ cần thiết.",
            "Nếu task lớn, leader tạo subtask.",
            "Member được giao nhìn thấy task trong danh sách của mình.",
            "Hệ thống có thể tạo thông báo khi task sắp đến hạn hoặc quá hạn.",
        ]),
        ("6.4. Thực hiện, báo cáo và review", [
            "Member mở task được giao và cập nhật trạng thái thực hiện.",
            "Member cập nhật phần trăm tiến độ, nội dung công việc và minh chứng.",
            "Member nộp báo cáo task.",
            "Leader xem báo cáo, attachment và lịch sử cập nhật.",
            "Leader ghi feedback, chấp nhận hoàn thành hoặc yêu cầu bổ sung.",
        ]),
        ("6.5. Theo dõi dashboard, lịch và thông báo", [
            "Leader mở dashboard sự kiện hoặc dashboard phòng ban để xem tiến độ.",
            "Leader lọc theo khoảng ngày, phòng ban, người thực hiện hoặc trạng thái.",
            "Leader tạo các mốc lịch quan trọng.",
            "Member xem lịch, nhận thông báo và đánh dấu thông báo đã đọc.",
        ]),
    ]
    for title, steps in flows:
        heading(doc, title, 2)
        for step in steps:
            number(doc, step)

    heading(doc, "7. Phân quyền người dùng", 1)
    paragraph(
        doc,
        "Phân quyền trong EventFlow được gán theo từng sự kiện, không phải vai trò toàn cục. Một người có thể là LEADER ở sự kiện A nhưng là MEMBER ở sự kiện B.",
    )
    permission_rows = [
        ("Xem danh sách sự kiện", "Không", "Không", "Sự kiện mình tham gia", "Sự kiện mình tham gia/quản lý"),
        ("Tạo sự kiện", "Không", "Có sau đăng nhập", "Có sau đăng nhập", "Có sau đăng nhập"),
        ("Xem chi tiết sự kiện", "Không", "Không", "Có", "Có"),
        ("Sửa/xóa sự kiện", "Không", "Không", "Không", "Có"),
        ("Xem/thêm/xóa thành viên", "Không", "Không", "Xem giới hạn", "Toàn quyền"),
        ("Tạo/sửa/xóa phòng ban", "Không", "Không", "Không", "Có"),
        ("Xem danh sách task", "Không", "Không", "Task được giao", "Toàn bộ task trong event"),
        ("Tạo/sửa/xóa task/subtask", "Không", "Không", "Không", "Có"),
        ("Cập nhật trạng thái task", "Không", "Không", "Task được giao", "Có"),
        ("Cập nhật tiến độ/nộp báo cáo", "Không", "Không", "Task được giao", "Có khi có quyền xem task"),
        ("Review task", "Không", "Không", "Không", "Có"),
        ("Xem dashboard", "Không", "Không", "Không", "Có"),
        ("Xem lịch sự kiện", "Không", "Không", "Có", "Có"),
        ("Tạo/sửa lịch sự kiện", "Không", "Không", "Không", "Có"),
        ("Upload attachment", "Không", "Không", "Task được giao", "Có"),
        ("Quản lý hồ sơ/thông báo", "Không", "Chỉ chính mình", "Chỉ chính mình", "Chỉ chính mình"),
    ]
    add_simple_table(
        doc,
        ["Nghiệp vụ", "Chưa đăng nhập", "Không thuộc event", "Member", "Leader"],
        permission_rows,
        [2400, 1500, 1650, 1800, 2010],
    )

    heading(doc, "8. Giá trị hệ thống mang lại", 1)
    heading(doc, "8.1. Giá trị cho leader", 2)
    for item in [
        "Có tầm nhìn tổng thể về toàn bộ sự kiện.",
        "Giảm thời gian tổng hợp tiến độ thủ công.",
        "Phát hiện sớm task trễ hạn, phòng ban chậm tiến độ hoặc người được giao quá tải.",
        "Quản lý rõ người chịu trách nhiệm, deadline và kết quả từng công việc.",
        "Có lịch sử báo cáo/review để đánh giá sau sự kiện.",
    ]:
        bullet(doc, item)
    heading(doc, "8.2. Giá trị cho member", 2)
    for item in [
        "Biết rõ việc cần làm, deadline và mức ưu tiên.",
        "Có một nơi chính thức để cập nhật tiến độ và nộp minh chứng.",
        "Nhận nhắc việc kịp thời qua ứng dụng, email hoặc Telegram.",
        "Dễ tìm lại tài liệu, feedback và lịch sử công việc.",
    ]:
        bullet(doc, item)
    heading(doc, "8.3. Giá trị cho tổ chức", 2)
    for item in [
        "Chuẩn hóa quy trình tổ chức sự kiện.",
        "Giảm thất lạc thông tin và giảm phụ thuộc vào nhóm chat.",
        "Tăng tính minh bạch trong phân công và đánh giá hiệu quả.",
        "Có dữ liệu để cải thiện quy trình cho các sự kiện sau.",
        "Có nền tảng kỹ thuật rõ ràng để mở rộng trong tương lai.",
    ]:
        bullet(doc, item)
    heading(doc, "8.4. Giá trị kỹ thuật", 2)
    for item in [
        "Kiến trúc tách frontend/backend rõ ràng, dễ phát triển độc lập.",
        "Có migration database bằng Flyway, phù hợp triển khai liên tục.",
        "Có Docker Compose để đóng gói môi trường.",
        "Có JWT, rate limiting, abuse protection và audit log để tăng độ an toàn.",
        "Có OpenAPI/Swagger và file API examples để kiểm thử.",
        "Có tùy chọn lưu file local hoặc Supabase Storage/S3.",
    ]:
        bullet(doc, item)

    heading(doc, "9. Kết luận", 1)
    paragraph(
        doc,
        "EventFlow là hệ thống quản lý điều phối sự kiện theo hướng thực tế, tập trung vào ba trụ cột: phân công rõ ràng, theo dõi tiến độ minh bạch và nhắc việc kịp thời. Hệ thống phù hợp với các sự kiện cần nhiều nhóm phối hợp, có leader điều phối và member thực hiện theo task.",
    )
    paragraph(
        doc,
        "Với các tính năng hiện có như quản lý sự kiện, phòng ban, thành viên, task/subtask, báo cáo, review, attachment, dashboard, lịch, thông báo, Telegram và AI assistant, EventFlow đáp ứng tốt nhu cầu MVP cho một nền tảng quản lý công việc tổ chức sự kiện nội bộ.",
    )
    paragraph(
        doc,
        "Trong các phiên bản tiếp theo, hệ thống có thể mở rộng thêm quản lý khách tham dự, đăng ký tham gia, check-in QR, bán vé/thanh toán, quản lý nhà tài trợ, ngân sách sự kiện và báo cáo sau sự kiện nâng cao. Ở phạm vi hiện tại, EventFlow đã đặt nền tảng vững chắc cho việc số hóa quy trình điều phối và giảm tải cho ban tổ chức.",
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
